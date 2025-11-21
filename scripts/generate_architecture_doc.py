"""
Generate Phoenix AI Architecture & Workflow Documentation (Word format)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sys

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    """Add a paragraph with optional style"""
    p = doc.add_paragraph(text, style=style)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_table_from_data(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
    
    return table

def create_architecture_document():
    """Generate the complete Word document"""
    doc = Document()
    
    # Title page
    title = doc.add_heading('Phoenix AI', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Technical Architecture & Workflow Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    version_info = doc.add_paragraph('Version: 1.0 | Date: November 2025')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== Section 1 ==========
    add_heading(doc, '1. High-Level System Overview', level=1)
    
    add_paragraph(doc, 'Phoenix AI unifies three intelligent testing domains:')
    doc.add_paragraph('PR & Code Impact Analysis', style='List Bullet')
    doc.add_paragraph('API Testing (Natural Language → Gherkin → Execution)', style='List Bullet')
    doc.add_paragraph('Autonomous Browser Interaction (UI validation & multi-step flows)', style='List Bullet')
    
    add_heading(doc, 'Core Building Blocks', level=2)
    building_blocks = [
        ['Component', 'Purpose'],
        ['WebUI (Gradio)', 'Collects user intent via tabbed interface'],
        ['Agent Settings', 'Supplies LLM provider, model, API key, base URL'],
        ['Domain Agents', 'Orchestrate workflows (PR, API, Browser)'],
        ['CustomBrowser', 'Wraps Playwright launch & context creation'],
        ['CustomController', 'Registers & executes actions (navigation, input, extraction)'],
        ['LLM', 'Drives planning (goals → structured actions)'],
        ['Artifacts', 'JSON results, GIF replay, screenshots, agent history'],
    ]
    add_table_from_data(doc, building_blocks[0], building_blocks[1:])
    
    doc.add_page_break()
    
    # ========== Section 2 ==========
    add_heading(doc, '2. Configuration & Initialization Flow', level=1)
    
    add_paragraph(doc, 'Step-by-step initialization process:')
    
    add_heading(doc, 'Step 1: WebUI Launch', level=3)
    add_paragraph(doc, 'User opens WebUI → Gradio loads components for Agent Settings, Browser Settings, API Testing, PR Testing, Browser Use tabs.')
    
    add_heading(doc, 'Step 2: Configuration Entry', level=3)
    add_paragraph(doc, 'User enters:')
    doc.add_paragraph('LLM provider (e.g., openai, google, anthropic, ollama)', style='List Bullet')
    doc.add_paragraph('Model name (e.g., gpt-4o, gemini-2.0-flash)', style='List Bullet')
    doc.add_paragraph('API key / base URL (optional for self-hosted endpoints)', style='List Bullet')
    
    add_paragraph(doc, 'These settings are stored in component registry (webui_manager.id_to_component), available to all tabs.')
    
    add_heading(doc, 'Step 3: Agent Invocation', level=3)
    add_paragraph(doc, 'When an agent is invoked:')
    doc.add_paragraph('Tab handler fetches current values from component registry', style='List Bullet')
    doc.add_paragraph('Calls llm_provider.get_llm_model() to initialize LLM object', style='List Bullet')
    doc.add_paragraph('Browser-specific settings passed into BrowserConfig', style='List Bullet')
    
    add_heading(doc, 'Step 4: Browser Initialization', level=3)
    add_paragraph(doc, 'On first browser task:')
    doc.add_paragraph('CustomBrowser launches Playwright with anti-detection flags', style='List Bullet')
    doc.add_paragraph('Custom window size, remote debugging port handling', style='List Bullet')
    doc.add_paragraph('CustomBrowserContext merges global and per-run overrides', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== Section 3 ==========
    add_heading(doc, '3. PR Testing Lifecycle', level=1)
    
    add_heading(doc, 'Workflow Steps', level=2)
    
    add_heading(doc, 'Step 1: PR Selection', level=3)
    add_paragraph(doc, 'User selects organization/project/repository → System pulls PR list via Azure DevOps client or analogous API.')
    
    add_heading(doc, 'Step 2: Data Fetching', level=3)
    add_paragraph(doc, 'PR Testing agent fetches:')
    doc.add_paragraph('PR metadata (title, description, commits)', style='List Bullet')
    doc.add_paragraph('Changed files and diffs', style='List Bullet')
    doc.add_paragraph('Associated work items (if integrated)', style='List Bullet')
    
    add_heading(doc, 'Step 3: Impact Analysis', level=3)
    add_paragraph(doc, 'Code analyzer scans diff to classify changes: modules, potential API endpoints, UI components.')
    
    add_heading(doc, 'Step 4: Test Plan Generation', level=3)
    add_paragraph(doc, 'LLM prompt constructed with:')
    doc.add_paragraph('Changed file paths & diff snippets', style='List Bullet')
    doc.add_paragraph('Work item context', style='List Bullet')
    doc.add_paragraph('Historical patterns (if memory enabled)', style='List Bullet')
    add_paragraph(doc, 'LLM outputs structured test plan: scenarios, manual steps, potential automated steps.')
    
    add_heading(doc, 'Step 5: Output & Artifacts', level=3)
    artifacts_table = [
        ['Artifact', 'Purpose'],
        ['Test plan JSON', 'Structured summary of PR metadata'],
        ['Test plan text', 'Human-readable scenarios'],
        ['Risk analysis logs', 'Indicate high-impact areas'],
    ]
    add_table_from_data(doc, artifacts_table[0], artifacts_table[1:])
    
    doc.add_page_break()
    
    # ========== Section 4 ==========
    add_heading(doc, '4. API Testing Lifecycle (Simplified Gherkin Flow)', level=1)
    
    add_heading(doc, 'Current Implementation', level=2)
    
    add_heading(doc, 'Step 1: Natural Language Input', level=3)
    add_paragraph(doc, 'User provides scenario description:')
    add_code_block(doc, 'Test GET https://jsonplaceholder.typicode.com/users and confirm Leanne Graham appears.')
    
    add_heading(doc, 'Step 2: Gherkin Generation', level=3)
    add_paragraph(doc, 'APITestingAgent.generate_gherkin_from_prompt():')
    doc.add_paragraph('Constructs system + user prompt for LLM', style='List Bullet')
    doc.add_paragraph('LLM returns structured Gherkin text', style='List Bullet')
    doc.add_paragraph('Agent formats into canonical Gherkin', style='List Bullet')
    
    add_paragraph(doc, 'Generated Gherkin example:')
    gherkin_example = '''Feature: Users API
Scenario: Validate users endpoint
  Given the API endpoint is "https://jsonplaceholder.typicode.com/users"
  When I send a GET request
  Then the response status code should be 200
  Then the response should be a JSON array
  Then the response should contain a user with name "Leanne Graham"'''
    add_code_block(doc, gherkin_example)
    
    add_heading(doc, 'Step 3: Execution', level=3)
    add_paragraph(doc, 'execute_gherkin_scenario() performs:')
    doc.add_paragraph('Parse URL (from Given line)', style='List Bullet')
    doc.add_paragraph('Parse HTTP method (from When line)', style='List Bullet')
    doc.add_paragraph('Extract expected status code (Then with "status code should be …")', style='List Bullet')
    doc.add_paragraph('Extract content assertions ("should contain" …)', style='List Bullet')
    doc.add_paragraph('Check array structure ("should be a JSON array")', style='List Bullet')
    
    add_heading(doc, 'Step 4: HTTP Request', level=3)
    add_paragraph(doc, 'HTTP request performed with requests.request(method, url, headers=..., timeout=...)')
    
    add_heading(doc, 'Step 5: Validation & Results', level=3)
    add_paragraph(doc, 'Response post-processing:')
    doc.add_paragraph('Attempt JSON parse, fallback to raw text', style='List Bullet')
    doc.add_paragraph('Validate status code', style='List Bullet')
    doc.add_paragraph('Search content for specified substrings', style='List Bullet')
    doc.add_paragraph('Record extracted boolean flags', style='List Bullet')
    
    add_paragraph(doc, 'Result JSON example:')
    result_json = '''{
  "test_status": "passed",
  "message": "Test passed successfully",
  "execution_time": "0.62s",
  "api_endpoint": "https://jsonplaceholder.typicode.com/users",
  "status_code": 200,
  "response_time_ms": 622.135,
  "full_response": [ ... entire JSON array ... ],
  "extracted_values": { "contains_Leanne_Graham": true },
  "validation_errors": [],
  "timestamp": "2025-11-20T18:34:35.769096"
}'''
    add_code_block(doc, result_json)
    
    doc.add_page_break()
    
    # ========== Section 5 ==========
    add_heading(doc, '5. Browser Autonomy Lifecycle', level=1)
    
    add_heading(doc, 'End-to-End Loop', level=2)
    
    add_heading(doc, 'Step 1: Task Input', level=3)
    add_paragraph(doc, 'User enters multi-step task or pastes Gherkin scenario into BrowserUse agent tab.')
    
    add_heading(doc, 'Step 2: Initialization', level=3)
    add_paragraph(doc, 'System performs:')
    doc.add_paragraph('LLM loaded (if credentials present)', style='List Bullet')
    doc.add_paragraph('CustomBrowser creates/gets context (viewport, proxy, headless)', style='List Bullet')
    doc.add_paragraph('CustomController registers actions (navigate, click, input, extract, upload)', style='List Bullet')
    
    add_heading(doc, 'Step 3: Planning Cycle', level=3)
    add_paragraph(doc, 'Agent builds current state snapshot:')
    doc.add_paragraph('Current URL', style='List Bullet')
    doc.add_paragraph('DOM summary', style='List Bullet')
    doc.add_paragraph('Memory excerpts from prior steps', style='List Bullet')
    
    add_paragraph(doc, 'LLM receives task goal, prior steps, extracted content, and outputs next action JSON.')
    
    add_heading(doc, 'Step 4: Action Dispatch', level=3)
    add_paragraph(doc, 'CustomController.act() process:')
    doc.add_paragraph('Iterates action JSON keys', style='List Bullet')
    doc.add_paragraph('Matches key to registered action in registry', style='List Bullet')
    doc.add_paragraph('Executes Playwright primitive (navigate, click, fill, extract)', style='List Bullet')
    
    action_table = [
        ['Action Type', 'Playwright Operation'],
        ['Navigation', 'page.goto()'],
        ['Click', 'element.click() via DOM index'],
        ['Input text', 'element.fill()'],
        ['Extraction', 'Fetch inner HTML or text'],
        ['Upload', 'set_input_files() on file input element'],
    ]
    add_table_from_data(doc, action_table[0], action_table[1:])
    
    add_heading(doc, 'Step 5: Post-Action Processing', level=3)
    add_paragraph(doc, 'After each action:')
    doc.add_paragraph('Screenshot captured (base64) added to step history', style='List Bullet')
    doc.add_paragraph('ActionResult includes extracted_content, error, include_in_memory', style='List Bullet')
    doc.add_paragraph('Memory summary updated (LLM can compress long contexts)', style='List Bullet')
    
    add_heading(doc, 'Step 6: Iteration', level=3)
    add_paragraph(doc, 'Loop continues until:')
    doc.add_paragraph('Explicit DONE action emitted', style='List Bullet')
    doc.add_paragraph('Failure repeated (could escalate to ask_for_assistant)', style='List Bullet')
    doc.add_paragraph('Max steps/time budget reached', style='List Bullet')
    
    add_heading(doc, 'Step 7: Finalization', level=3)
    add_paragraph(doc, 'On completion:')
    doc.add_paragraph('GIF assembled from step screenshots', style='List Bullet')
    doc.add_paragraph('Agent history JSON saved under tmp/agent_history/<uuid>/', style='List Bullet')
    doc.add_paragraph('WebUI shows task result + memory summary + screenshot progression', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== Section 6 ==========
    add_heading(doc, '6. Element Handling & Robustness', level=1)
    
    add_heading(doc, 'Mechanisms', level=2)
    
    add_heading(doc, 'DOM Indexing', level=3)
    add_paragraph(doc, 'browser_use maps visible DOM elements to indices for stable referencing. Actions target elements using numeric index, avoiding brittle CSS selectors.')
    
    add_heading(doc, 'Locator Resolution', level=3)
    add_paragraph(doc, 'For file upload and interaction: browser.get_dom_element_by_index(index) returns structured DOM wrapper with locate functions for Playwright locator access.')
    
    add_heading(doc, 'Anti-Detection', level=3)
    add_paragraph(doc, 'Chrome launched with selective args:')
    doc.add_paragraph('Remove automation flags', style='List Bullet')
    doc.add_paragraph('Custom window position & size for realism', style='List Bullet')
    
    add_heading(doc, 'Fallback & Errors', level=3)
    add_paragraph(doc, 'If element not found: ActionResult(error="...") returned; LLM receives error and re-plans.')
    add_paragraph(doc, 'Upload file path validation performed prior to action.')
    
    add_heading(doc, 'Content Extraction', level=3)
    add_paragraph(doc, 'ExtractPageContentAction retrieves structured JSON or raw text for validation and memory enhancement.')
    
    add_heading(doc, 'Memory Inclusion', level=3)
    add_paragraph(doc, 'Actions may set include_in_memory=True so extracted text appears in subsequent planning context.')
    
    doc.add_page_break()
    
    # ========== Section 7 ==========
    add_heading(doc, '7. Data & Observability', level=1)
    
    add_heading(doc, 'Collected Artifacts', level=2)
    
    artifacts_comprehensive = [
        ['Artifact', 'Purpose', 'Location'],
        ['Test Result JSON', 'API test outcome', 'In-memory, displayed in UI'],
        ['Full API Response', 'Debugging, evidence', 'Embedded in result JSON'],
        ['Agent History JSON', 'Replay & audit', 'tmp/agent_history/<uuid>/<uuid>.json'],
        ['GIF Session Replay', 'Visual narrative', 'Same folder as history JSON'],
        ['Screenshots (base64)', 'Step-by-step state', 'Stored inside history JSON'],
        ['Logs', 'Root cause analysis', 'Console / future file sink'],
        ['Gherkin Scenario Text', 'Reproducible scenario', 'UI textbox / user editable'],
    ]
    add_table_from_data(doc, artifacts_comprehensive[0], artifacts_comprehensive[1:])
    
    add_heading(doc, 'Observability Principles', level=2)
    doc.add_paragraph('Every meaningful action produces a structured record', style='List Bullet')
    doc.add_paragraph('Human-reviewable evidence (GIF + JSON) supports trust & reproducibility', style='List Bullet')
    doc.add_paragraph('Failures annotated with error content for iterative AI correction', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== Section 8 ==========
    add_heading(doc, '8. Error Handling & Recovery Strategy', level=1)
    
    error_categories = [
        ['Error Category', 'Handling Strategy'],
        ['LLM Init Failure', 'UI warns user to supply API key; agent proceeds with None for limited tasks'],
        ['Network/HTTP errors', 'requests.RequestException → test result flagged with error status'],
        ['Element Not Found', 'Action-level error → LLM receives context → may scroll/search/escalate'],
        ['Gherkin Parse Issues', 'Missing Feature:/Scenario: triggers validation error; user corrects'],
        ['Browser Launch Conflict', 'Remote debugging port conflict resolved by removing port arg'],
        ['Graceful Abort', 'Repeated errors trigger ask_for_assistant action (future enhancement)'],
    ]
    add_table_from_data(doc, error_categories[0], error_categories[1:])
    
    doc.add_page_break()
    
    # ========== Section 9 ==========
    add_heading(doc, '9. Extensibility Points', level=1)
    
    extensibility = [
        ['Extension Point', 'How to Extend'],
        ['New Actions', 'Add decorator to CustomController._register_custom_actions()'],
        ['MCP Tools', 'Start MCP client → dynamic registration of external server tools'],
        ['API Assertions', 'Enhance execute_gherkin_scenario() with schema validator (jsonschema)'],
        ['Bulk Test Generation', 'Add OpenAPI importer → map endpoints → generate multi-scenario Gherkin'],
        ['CI Integration', 'Serialize test results, add CLI runner for regression batch'],
        ['Visual Validation', 'Extend BrowserUseAgent for DOM diffs / screenshot comparisons'],
        ['Auth Flows', 'Augment Gherkin parser for tokens, sessions, reuse cookies across contexts'],
    ]
    add_table_from_data(doc, extensibility[0], extensibility[1:])
    
    doc.add_page_break()
    
    # ========== Section 10 ==========
    add_heading(doc, '10. End-to-End Example Walkthrough', level=1)
    
    add_heading(doc, 'Scenario', level=2)
    add_paragraph(doc, '"Validate users API and confirm UI shows user \'Leanne Graham\'."')
    
    add_heading(doc, 'Step-by-Step Execution', level=2)
    
    doc.add_paragraph('User enters API description → Gherkin generated', style='List Number')
    doc.add_paragraph('Execute API test → Response JSON captured (contains user list)', style='List Number')
    doc.add_paragraph('User copies summary or reuses scenario in BrowserUse tab', style='List Number')
    doc.add_paragraph('BrowserUse Agent plan:', style='List Number')
    
    p = doc.add_paragraph('Navigate to API URL', style='List Bullet 2')
    p = doc.add_paragraph('Extract content, confirm Leanne Graham present', style='List Bullet 2')
    p = doc.add_paragraph('Done action once validations satisfied', style='List Bullet 2')
    
    doc.add_paragraph('Artifacts collected:', style='List Number')
    p = doc.add_paragraph('API Test JSON Result', style='List Bullet 2')
    p = doc.add_paragraph('Browser GIF & History JSON', style='List Bullet 2')
    p = doc.add_paragraph('Combined narrative for presentation', style='List Bullet 2')
    
    doc.add_page_break()
    
    # ========== Section 11 ==========
    add_heading(doc, '11. Architecture Diagrams', level=1)
    
    add_heading(doc, 'API Testing Flow', level=2)
    api_flow = '''User Prompt
   │
   ▼
APITestingAgent.generate_gherkin_from_prompt()
   │  (LLM transforms intent → structured scenario)
   ▼
Gherkin Text (Feature/Scenario/Given/When/Then)
   │
   ▼
execute_gherkin_scenario()
   │  Parse URL/method/status/assertions
   │  Perform HTTP request
   │  Validate + assemble full_response
   ▼
Result JSON (test_status, full_response, extracted_values, timing)'''
    add_code_block(doc, api_flow)
    
    add_heading(doc, 'Browser Autonomy Flow', level=2)
    browser_flow = '''User Task / Scenario
   │
   ▼
LLM Planning Loop
   │  (Goal reasoning, memory summarization)
   ▼
Structured Action JSON
   │
   ▼
CustomController.act()
   │  (Lookup registered actions)
   ▼
Playwright Execution (navigation/input/extract/upload)
   │
   ├─ Screenshot capture
   ├─ ActionResult with extracted content
   └─ Update memory/context
   │
   └─ Repeat until DONE or failure
   ▼
Artifacts (GIF, history JSON, logs)'''
    add_code_block(doc, browser_flow)
    
    doc.add_page_break()
    
    # ========== Section 12 ==========
    add_heading(doc, '12. Key Design Principles', level=1)
    
    doc.add_paragraph('Declarative Prompts → Deterministic Structured Outputs', style='List Bullet')
    doc.add_paragraph('Separation of concerns (Agent orchestration vs. execution engines)', style='List Bullet')
    doc.add_paragraph('Resilience & transparency (Artifacts before sophistication)', style='List Bullet')
    doc.add_paragraph('Extensibility over monolithic complexity (registry patterns)', style='List Bullet')
    doc.add_paragraph('Minimal configuration friction (derive from Gherkin; central settings)', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== Section 13 ==========
    add_heading(doc, '13. Future Roadmap', level=1)
    
    roadmap = [
        ['Feature', 'Target', 'Description'],
        ['Schema Validation', 'Q1 2025', 'JSON Schema enforcement for API responses'],
        ['OpenAPI Import', 'Q1 2025', 'Bulk Gherkin generation from API specs'],
        ['Performance Budgets', 'Q2 2025', 'Response time thresholds & alerts'],
        ['Flakiness Tracking', 'Q2 2025', 'Stability scoring for test reliability'],
        ['CI/CD Integration', 'Q2 2025', 'Export test results for pipeline gating'],
        ['Visual Regression', 'Q3 2025', 'Screenshot comparison & DOM diff analysis'],
        ['Knowledge Graph', 'Q4 2025', 'Link code changes to docs automatically'],
    ]
    add_table_from_data(doc, roadmap[0], roadmap[1:])
    
    doc.add_page_break()
    
    # ========== Appendix ==========
    add_heading(doc, 'Appendix A: Glossary', level=1)
    
    glossary = [
        ['Term', 'Definition'],
        ['Gherkin', 'Structured test description format (Given/When/Then)'],
        ['LLM', 'Large Language Model (AI reasoning engine)'],
        ['Endpoint', 'API URL for HTTP operations'],
        ['Action', 'Single browser operation (Click, Navigate, etc.)'],
        ['Artifact', 'Saved evidence file (JSON, GIF, screenshot)'],
        ['DOM Index', 'Numeric reference to page elements'],
        ['ActionResult', 'Structured return object from actions'],
        ['CustomController', 'Action registry and dispatcher'],
        ['CustomBrowser', 'Playwright wrapper with configuration'],
    ]
    add_table_from_data(doc, glossary[0], glossary[1:])
    
    doc.add_page_break()
    
    add_heading(doc, 'Appendix B: Configuration Reference', level=1)
    
    config_ref = [
        ['Setting', 'Purpose', 'Example'],
        ['LLM Provider', 'AI service selection', 'openai, anthropic, google'],
        ['Model Name', 'Specific model version', 'gpt-4o, gemini-2.0-flash'],
        ['API Key', 'Authentication credential', 'sk-...'],
        ['Base URL', 'Custom endpoint (optional)', 'https://api.custom.com'],
        ['Headless', 'Browser visibility', 'true/false'],
        ['Viewport', 'Window dimensions', '1280x1100'],
        ['User Data Dir', 'Browser profile path', 'C:/Users/.../Chrome/User Data'],
    ]
    add_table_from_data(doc, config_ref[0], config_ref[1:])
    
    # Save document
    output_path = 'Phoenix_AI_Architecture_Documentation.docx'
    doc.save(output_path)
    print(f"✅ Document created successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        output = create_architecture_document()
        print(f"\n📄 Architecture documentation saved to: {output}")
        print("\nDocument includes:")
        print("  • High-level system overview")
        print("  • Configuration & initialization flows")
        print("  • PR Testing lifecycle")
        print("  • API Testing (Gherkin) workflows")
        print("  • Browser autonomy execution loops")
        print("  • Element handling & robustness strategies")
        print("  • Data & observability artifacts")
        print("  • Error handling & recovery")
        print("  • Extensibility points")
        print("  • End-to-end examples")
        print("  • Architecture diagrams")
        print("  • Design principles")
        print("  • Future roadmap")
        print("  • Glossary & configuration reference")
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
